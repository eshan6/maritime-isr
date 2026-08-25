"""Bytes to passages — the only place that knows what a PDF is.

Every reader returns the same thing: a list of :class:`Passage`, each carrying
the text, a locator an operator can find **by eye**, and the confidence the
reader earned for it. Nothing downstream knows or cares which format a
notification arrived in, which is what lets the electronic feed enter at the
same seam as a scanned fax.

**A locator is not an offset.** "page 1", "Sheet1!B7", "table 2 row 3" — an
analyst has to be able to open the attachment and put a finger on the line. A
character range satisfies a schema and helps nobody.

**Confidence is per-passage and it is earned, not assumed.** A spreadsheet cell
is 1.0 because there is no ambiguity about what was in it. A PDF text layer is
0.97 because layout can interleave columns. OCR is whatever tesseract reports,
floored — and a fax at 0.55 has to be *visibly* less trustworthy than a portal
field at 1.0, or the whole per-field provenance discipline collapses into
decoration.
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

__all__ = ["Passage", "READERS", "read_document", "reader_availability",
           "ReaderUnavailable", "OCR_MIN_CONFIDENCE"]


#: Floor for an OCR'd passage. Tesseract's per-word confidences average out
#: optimistically on a clean render, and the point of the floor is that no
#: OCR'd value ever reads as certain — a scan is the least trustworthy thing in
#: the inbox and the number has to say so even on a good day.
OCR_MIN_CONFIDENCE = 0.35
OCR_MAX_CONFIDENCE = 0.85


class ReaderUnavailable(RuntimeError):
    """The library or binary this format needs is not installed.

    Raised rather than returning empty, because an unreadable document and an
    empty document are different facts and the difference matters: one is a
    missing dependency, the other is a form somebody left blank. Collapsing
    them produces a notification with no fields and no explanation.
    """


@dataclass(frozen=True)
class Passage:
    """One readable chunk of a document, with where it came from."""

    text: str
    locator: str
    confidence: float
    method: str


# --------------------------------------------------------------------------
# PDF — text layer first, OCR only when there is no text
# --------------------------------------------------------------------------

def _pdf_text(path: Path) -> list[Passage]:
    try:
        from pypdf import PdfReader
    except ImportError as e:                                 # pragma: no cover
        raise ReaderUnavailable(f"pypdf not installed: {e}") from e

    out: list[Passage] = []
    reader = PdfReader(str(path))
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for line in text.splitlines():
            line = line.strip()
            if line:
                out.append(Passage(line, f"page {i}", 0.97, "pdf_text"))
    return out


def _pdf_ocr(path: Path) -> list[Passage]:
    """Render each page and read the pixels.

    Only reached when the text layer is empty, which is exactly what a scan or
    a fax looks like. Rendering is done with the PDF's own embedded images
    where possible — a scanned PDF is a wrapper around a picture, so the
    picture is what we want and re-rasterising the page would be a lossy
    round trip through a renderer this project does not ship.
    """
    try:
        import pytesseract
        from PIL import Image
        from pypdf import PdfReader
    except ImportError as e:
        raise ReaderUnavailable(f"OCR stack not installed: {e}") from e

    out: list[Passage] = []
    reader = PdfReader(str(path))
    for i, page in enumerate(reader.pages, start=1):
        for image_file in getattr(page, "images", []) or []:
            import io
            img = Image.open(io.BytesIO(image_file.data))
            try:
                data = pytesseract.image_to_data(
                    img, output_type=pytesseract.Output.DICT)
            except Exception as e:                           # noqa: BLE001
                raise ReaderUnavailable(
                    f"tesseract binary unavailable: {e}") from e
            out.extend(_passages_from_ocr(data, f"page {i} (scanned)"))
    return out


def _passages_from_ocr(data: dict, locator: str) -> list[Passage]:
    """Group tesseract's word boxes back into lines, keeping their confidence.

    Line grouping matters because the extractor reads `Label: value` pairs, and
    tesseract returns words. A word-per-passage stream would separate every
    label from its value and the reader would find nothing — which would look
    like OCR failing when it had in fact worked.
    """
    lines: dict[tuple, list[tuple[str, float]]] = {}
    n = len(data.get("text", []))
    for i in range(n):
        word = (data["text"][i] or "").strip()
        if not word:
            continue
        try:
            conf = float(data["conf"][i])
        except (TypeError, ValueError):
            conf = -1.0
        if conf < 0:
            continue
        key = (data["page_num"][i], data["block_num"][i],
               data["par_num"][i], data["line_num"][i])
        lines.setdefault(key, []).append((word, conf / 100.0))

    out = []
    for key in sorted(lines):
        words = lines[key]
        text = " ".join(w for w, _ in words)
        mean = sum(c for _, c in words) / len(words)
        conf = max(OCR_MIN_CONFIDENCE, min(OCR_MAX_CONFIDENCE, mean))
        out.append(Passage(text, locator, round(conf, 3), "ocr"))
    return out


def read_pdf(path: Path) -> list[Passage]:
    """Text layer if there is one, pixels if there is not.

    The order is not an optimisation. A PDF with a text layer is authoritative
    about its own content; OCR of the same page is a guess at it. Trying OCR
    first would replace certainty with an estimate on every document that did
    not need it.
    """
    passages = _pdf_text(path)
    if passages:
        return passages
    return _pdf_ocr(path)


# --------------------------------------------------------------------------
# Word — paragraphs *and* tables
# --------------------------------------------------------------------------

def read_docx(path: Path) -> list[Passage]:
    """Both shapes, because both are common and each is silent about the other.

    A reader that walks paragraphs finds nothing in a table-shaped form, and
    the failure looks like an empty submission rather than an unread one. Table
    cells are joined with a colon so the extractor sees the same `Label: value`
    shape it sees everywhere else — one grammar, four formats.
    """
    try:
        from docx import Document
    except ImportError as e:                                 # pragma: no cover
        raise ReaderUnavailable(f"python-docx not installed: {e}") from e

    doc = Document(str(path))
    out: list[Passage] = []
    for i, para in enumerate(doc.paragraphs, start=1):
        text = (para.text or "").strip()
        if text:
            out.append(Passage(text, f"paragraph {i}", 1.0, "docx_paragraph"))
    for t, table in enumerate(doc.tables, start=1):
        for r, row in enumerate(table.rows, start=1):
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if len(cells) >= 2:
                out.append(Passage(f"{cells[0]}: {cells[1]}",
                                   f"table {t} row {r}", 1.0, "docx_table"))
            elif cells:
                out.append(Passage(cells[0], f"table {t} row {r}", 1.0,
                                   "docx_table"))
    return out


# --------------------------------------------------------------------------
# spreadsheet
# --------------------------------------------------------------------------

def read_xlsx(path: Path) -> list[Passage]:
    """Label and value as adjacent cells, wherever on the sheet they start.

    Real submitted spreadsheets have a logo, a blank row and a heading before
    the form begins, so the label column is not column A. Scanning every row for
    a cell followed by a non-empty neighbour finds the form wherever it starts,
    and costs nothing on a sheet this size.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as e:                                 # pragma: no cover
        raise ReaderUnavailable(f"openpyxl not installed: {e}") from e

    wb = load_workbook(str(path), data_only=True, read_only=True)
    out: list[Passage] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            cells = [(c, c.value) for c in row if c.value not in (None, "")]
            for i in range(len(cells) - 1):
                label_cell, label = cells[i]
                value_cell, value = cells[i + 1]
                if label_cell.row != value_cell.row:
                    continue
                out.append(Passage(
                    f"{label}: {value}",
                    f"{ws.title}!{label_cell.coordinate}", 1.0, "xlsx_cell"))
            if len(cells) == 1:
                cell, value = cells[0]
                out.append(Passage(str(value),
                                   f"{ws.title}!{cell.coordinate}", 1.0,
                                   "xlsx_cell"))
    wb.close()
    return out


# --------------------------------------------------------------------------
# the electronic feed — a reader, not a pipeline
# --------------------------------------------------------------------------

#: Portal JSON paths to the canonical field names. The portal's shape is not
#: ours and never will be; mapping it here is the whole adaptation, and it is
#: the reason the requirement's "without rework" clause holds.
_ELECTRONIC_MAP = {
    "submittedAt": "filed_at",
    "vessel.name": "vessel_name",
    "vessel.imo": "imo",
    "vessel.callSign": "call_sign",
    "vessel.flag": "flag",
    "voyage.lastPort": "last_port",
    "voyage.arrivalPort": "arrival_port",
    "voyage.eta": "eta",
    "voyage.cargo": "cargo",
    "parties.owner": "owner",
    "parties.agent": "agent",
    "parties.crew": "crew_count",
}


def read_electronic(path: Path) -> list[Passage]:
    """The portal feed, flattened to the same `Label: value` grammar.

    A structured feed could bypass the extractor entirely and set fields
    directly. It deliberately does not. Making it produce passages means the
    electronic path and the fax path are exercised by the *same* extraction
    code, so a change to how a date is parsed cannot fix one and break the
    other — and the "drops in without rework" claim is demonstrated by the
    shared code rather than asserted in a design note.
    """
    payload = json.loads(Path(path).read_text(encoding="utf-8"))
    out: list[Passage] = []
    for dotted, field_name in _ELECTRONIC_MAP.items():
        if "." in dotted:
            group, key = dotted.split(".")
            value = (payload.get(group) or {}).get(key)
        else:
            value = payload.get(dotted)
        if value in (None, ""):
            continue
        out.append(Passage(f"{field_name}: {value}", f"${dotted}", 1.0,
                           "electronic_field"))
    return out


READERS: dict[str, Callable[[Path], list[Passage]]] = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".xlsx": read_xlsx,
    ".json": read_electronic,
}


def read_document(path) -> list[Passage]:
    """Dispatch on extension. Unknown extensions are an explicit refusal."""
    path = Path(path)
    reader = READERS.get(path.suffix.lower())
    if reader is None:
        raise ReaderUnavailable(
            f"no reader for {path.suffix!r} — an inbox may hold anything, and "
            f"a format nobody wrote a reader for is a gap to report, not a row "
            f"to invent")
    return reader(path)


def reader_availability() -> dict:
    """Which formats this machine can actually read, and why not where not.

    Printed by the connector before it starts. A run that silently reads three
    of four formats reports a smaller inbox than arrived, and the missing
    quarter looks like nobody submitted rather than like nobody could read.
    """
    out: dict[str, str] = {}
    for name, module in (("pdf", "pypdf"), ("docx", "docx"),
                         ("xlsx", "openpyxl"), ("electronic", "json")):
        try:
            __import__(module)
            out[name] = "ok"
        except ImportError as e:
            out[name] = f"unavailable: {e}"
    # OCR is the one that can be installed and still not work, because
    # `pytesseract` is a wrapper around a binary pip cannot install.
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        out["pdf_scan"] = "ok"
    except ImportError as e:
        out["pdf_scan"] = f"unavailable: {e}"
    except Exception as e:                                   # noqa: BLE001
        out["pdf_scan"] = f"unavailable: tesseract binary missing ({e})"
    return out


def normalise_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()
