"""Pre-Arrival Notifications, in the formats they actually arrive in — Area 4.

*"Pre-Arrival Notification of Ships data reaches the Coast Guard as PDF, Word
and spreadsheet attachments by email. It contains vital information but cannot
be stored in a structured database or fused with AIS because of its format."*
— the IDEX Challenge 82 brief, Area 4.

And the instruction that governs everything in this module:

> Generate realistic notification documents from your existing scenario truth,
> in the messy formats the requirement names — including the ones that are
> scans, the ones with inconsistent field ordering, and the ones where a human
> typed the vessel name slightly differently than the registry spells it. **A
> clean document set proves nothing, because the entire difficulty here is that
> the input is unstructured.**

So the mess here is authored, deliberately, and each kind of it is a thing that
happens in a real inbox:

* **Four formats.** A text PDF, a scanned PDF that is an image of a page and
  carries no text layer at all, a Word document, a spreadsheet. Plus a fifth
  that is not a document: the structured electronic feed the national logistics
  portal is expected to publish, which the requirement asks the system to stay
  compatible with. It is generated here beside the others precisely so that
  compatibility is demonstrated rather than asserted.
* **Inconsistent field order.** Agencies use their own forms. Half the
  documents here put the last port before the ETA and half after, and one puts
  the cargo at the top because that is what its shipper cares about.
* **Inconsistent field *labels*.** "Last Port of Call", "From", "Previous
  Port", "Port of Departure" all mean one thing. A reader keyed to one spelling
  reads a third of a real corpus.
* **Names typed by hand.** "M.V. GRANITE TRIUMPH", "GRANITE TRIUMPH", "GRANITE
  TRUIMPH" — a prefix, a punctuation choice and a transposition, which is what
  a form filled in at 0300 by an agent's clerk looks like.
* **Missing fields.** Real notifications arrive incomplete. A crew count is
  absent about as often as it is present, and a reader that assumes every field
  is there will read a null as a value.

**What is NOT deliberately corrupted: the truth underneath.** These documents
describe vessels the corpus already contains, and the contradictions between
paperwork and track are authored as *scenarios* (group P), not sprinkled as
noise. A document that is merely garbled is a parsing problem; a document that
says something the track disagrees with is the product.
"""
from __future__ import annotations

import hashlib
import io
import json
import random
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

__all__ = ["NotificationSpec", "write_notifications", "FORMATS",
           "LABEL_SYNONYMS", "PANS_DIRNAME", "DOCUMENTS_DIRNAME",
           "DocumentKind", "DOCUMENT_KINDS", "HouseStyle", "HOUSE_STYLES",
           "PORT_AUTHORITIES", "build_document_specs", "AUTHORED_CASES",
           "CASE_MIX"]

#: Where the generated documents land, under the data root. They are *inputs*,
#: not conformed rows — the same standing an unread email attachment has — so
#: they sit beside `conformed/` rather than inside it.
PANS_DIRNAME = "pans_inbox"

#: Where the *wider* port-paperwork corpus lands — the arrival and departure
#: reports, crew lists, cargo manifests and clearance certificates that arrive
#: in the same mailbox as a PANS.
#:
#: **A second directory, not a second pipeline.** It exists so the two corpora
#: can be counted separately (the PANS inbox is what ADR-036's figures were
#: measured on, and mixing a second set into it would silently move them), and
#: because `scenario.run.clear` only knows how to empty `pans_inbox/`. Both are
#: read by the same connector, through the same readers, into the same table.
DOCUMENTS_DIRNAME = "port_documents"

#: The formats a notification arrives in. `electronic` is the portal feed the
#: requirement names as the future state, and it is generated here so the claim
#: "the electronic feed drops in without rework" can be tested rather than made.
FORMATS = ("pdf", "pdf_scan", "docx", "xlsx", "electronic")

#: What different agencies call the same field. A reader keyed to one spelling
#: reads a fraction of a real corpus, which is why the extractor is written
#: against this table and why the table lives here — the generator and the
#: reader must disagree about labels the way two agencies do, but they must not
#: disagree about which labels *exist*.
LABEL_SYNONYMS: dict[str, tuple[str, ...]] = {
    "vessel_name": ("Vessel Name", "Name of Vessel", "Ship Name", "Vessel"),
    "imo": ("IMO Number", "IMO No.", "IMO", "IMO/Official No."),
    "call_sign": ("Call Sign", "Callsign", "Radio Call Sign", "C/S"),
    "flag": ("Flag", "Flag State", "Nationality", "Flag of Registry"),
    "last_port": ("Last Port of Call", "From", "Previous Port",
                  "Port of Departure", "Last Port"),
    "arrival_port": ("Port of Arrival", "To", "Destination Port",
                     "Arrival Port", "Port"),
    "eta": ("ETA", "Expected Time of Arrival", "Estimated Arrival",
            "Date/Time of Arrival"),
    "cargo": ("Cargo", "Nature of Cargo", "Cargo on Board", "Commodity"),
    "crew_count": ("Crew", "No. of Crew", "Crew on Board", "Total Crew"),
    "owner": ("Owner", "Registered Owner", "Owners", "Beneficial Owner"),
    "agent": ("Agent", "Local Agent", "Ship's Agent", "Agency"),
    "filed_at": ("Date of Filing", "Filed On", "Date of Submission",
                 "Submitted", "Notification Date"),
}

#: Field orders real forms use. Not a shuffle — each is a plausible form layout,
#: because a random permutation per document would be a harder problem than the
#: real one and would test the reader against noise rather than against variety.
#:
#: `filed_at` sits where a form puts it — at the top on some layouts, in the
#: sign-off block at the bottom on others. It is a declared field like any
#: other and is read like one; see `_lines` for why that matters.
FIELD_ORDERS: tuple[tuple[str, ...], ...] = (
    ("vessel_name", "imo", "call_sign", "flag", "last_port", "arrival_port",
     "eta", "cargo", "crew_count", "owner", "agent", "filed_at"),
    ("filed_at", "vessel_name", "imo", "flag", "arrival_port", "eta",
     "last_port", "cargo", "agent", "owner", "crew_count", "call_sign"),
    ("cargo", "vessel_name", "imo", "last_port", "arrival_port", "eta",
     "owner", "agent", "flag", "call_sign", "crew_count", "filed_at"),
    ("filed_at", "imo", "vessel_name", "eta", "arrival_port", "last_port",
     "flag", "crew_count", "cargo", "owner", "agent", "call_sign"),
)


# --------------------------------------------------------------------------
# the letterheads — the ports these documents are actually filed at
# --------------------------------------------------------------------------

#: Port authority and address, by the gazetteer name of the port.
#:
#: **The letterhead names the authority; the declared value names the port.**
#: A form printed by Deendayal Port Authority says "Kandla" in its Port of
#: Arrival box, because that is the name the gazetteer and the AIS destination
#: field both use. Putting "Deendayal" in the *value* would resolve to nothing
#: and every rule reading it would answer "not checkable" — a corpus-wide
#: silence caused by the generator, which is the ADR-036 defect that cost most.
PORT_AUTHORITIES: dict[str, tuple[str, str]] = {
    "Kandla": ("DEENDAYAL PORT AUTHORITY, KANDLA",
               "Administrative Office, Gandhidham, Kachchh, Gujarat 370201"),
    "Mundra": ("MUNDRA PORT — OFFICE OF THE HARBOUR MASTER",
               "Mundra, Kachchh, Gujarat 370421"),
    "JNPT": ("JAWAHARLAL NEHRU PORT AUTHORITY — NHAVA SHEVA",
             "Administration Building, Sheva, Navi Mumbai 400707"),
    "Mumbai": ("MUMBAI PORT AUTHORITY — HARBOUR MASTER'S OFFICE",
               "Port House, Shoorji Vallabhdas Marg, Mumbai 400001"),
    "Mangalore": ("NEW MANGALORE PORT AUTHORITY",
                  "Panambur, Mangaluru, Karnataka 575010"),
    "Kochi": ("COCHIN PORT AUTHORITY",
              "Willingdon Island, Kochi, Kerala 682009"),
}

#: For a port with no authority of its own in the table above. Every one of
#: these documents is still filed *somewhere*, and a form with no letterhead is
#: not a form anybody in this region receives.
DEFAULT_AUTHORITY = (
    "OFFICE OF THE PRINCIPAL OFFICER, MERCANTILE MARINE DEPARTMENT",
    "Directorate General of Shipping, Government of India")


def authority_for(port) -> tuple[str, str]:
    """The letterhead a document filed at this port carries."""
    return PORT_AUTHORITIES.get(str(port or ""), DEFAULT_AUTHORITY)


# --------------------------------------------------------------------------
# the kinds of paper that arrive in one mailbox
# --------------------------------------------------------------------------

@dataclass(frozen=True)
class DocumentKind:
    """One kind of port paperwork: its title, and which fields it carries.

    **A kind is not a format.** A crew list arrives as a PDF, a Word table, a
    spreadsheet or a portal payload exactly as a PANS does, and it is read by
    the same four readers. What differs is the *title block* and which of the
    canonical fields the form has boxes for — which is why the arrival-window
    rule is `not_checkable` on a departure report and always will be: a
    departure report has no ETA box to fill in.
    """

    key: str
    title: str
    form_ref: str
    prefix: str
    #: Canonical fields this form carries boxes for, in the order a form of
    #: this kind puts them.
    fields: tuple[str, ...]
    #: Which body follows the header block: a crew table, a manifest, a
    #: certificate paragraph, or nothing.
    body: Optional[str] = None
    #: One line describing the form, printed under the title.
    note: str = ""


_IDENTITY = ("vessel_name", "imo", "call_sign", "flag")

DOCUMENT_KINDS: dict[str, DocumentKind] = {
    "pans": DocumentKind(
        key="pans",
        title="PRE-ARRIVAL NOTIFICATION OF SECURITY (PANS)",
        form_ref="Form PANS-1",
        prefix="PANS",
        fields=_IDENTITY + ("last_port", "arrival_port", "eta", "cargo",
                            "crew_count", "owner", "agent", "filed_at"),
        note="Submitted under the Merchant Shipping (Security) Rules, "
             "96 hours before entry into Indian waters."),
    "arrival_report": DocumentKind(
        key="arrival_report",
        title="ARRIVAL REPORT",
        form_ref="IMO FAL Form 1 (General Declaration — Arrival)",
        prefix="ARR",
        fields=_IDENTITY + ("last_port", "arrival_port", "eta", "cargo",
                            "crew_count", "agent", "filed_at"),
        note="Rendered by the Master or his Agent on arrival at the port."),
    "departure_report": DocumentKind(
        key="departure_report",
        title="DEPARTURE REPORT",
        form_ref="IMO FAL Form 1 (General Declaration — Departure)",
        prefix="DEP",
        # **No arrival port and no ETA, deliberately.** A departure report
        # declares where she has been and what she is carrying away; it has no
        # box for when she will next arrive. Every arrival-window check on one
        # of these answers `not_checkable`, and that is the correct answer
        # rather than a gap in the reader.
        fields=_IDENTITY + ("last_port", "cargo", "crew_count", "owner",
                            "agent", "filed_at"),
        note="Rendered before the vessel is granted clearance to proceed."),
    "crew_list": DocumentKind(
        key="crew_list",
        title="CREW LIST",
        form_ref="IMO FAL Form 5",
        prefix="CREW",
        fields=_IDENTITY + ("last_port", "arrival_port", "eta", "crew_count",
                            "agent", "filed_at"),
        body="crew",
        note="One line per person on board, including the Master."),
    "cargo_manifest": DocumentKind(
        key="cargo_manifest",
        title="CARGO DECLARATION (MANIFEST)",
        form_ref="IMO FAL Form 2",
        prefix="MFST",
        fields=_IDENTITY + ("last_port", "arrival_port", "eta", "cargo",
                            "agent", "filed_at"),
        body="manifest",
        note="Particulars of cargo carried, for the port of discharge shown."),
    "port_clearance": DocumentKind(
        key="port_clearance",
        title="PORT CLEARANCE CERTIFICATE",
        form_ref="Form MS-CLR",
        prefix="CLR",
        # No ETA on a clearance either: it certifies something that has already
        # happened.
        fields=_IDENTITY + ("arrival_port", "cargo", "crew_count", "owner",
                            "agent", "filed_at"),
        body="clearance",
        note="Granted under section 96 of the Merchant Shipping Act, 1958."),
}

#: The order kinds are cycled through, so every kind gets a comparable sample
#: whatever the corpus size. Same reasoning as the round-robin over formats in
#: `group_p.build_notifications`: a random draw at this scale leaves one kind
#: with three documents and makes any per-kind figure noise.
DOCUMENT_KIND_ORDER: tuple[str, ...] = (
    "pans", "arrival_report", "crew_list", "cargo_manifest",
    "departure_report", "port_clearance")


# --------------------------------------------------------------------------
# house styles — six agencies, six ways of asking the same twelve questions
# --------------------------------------------------------------------------

@dataclass(frozen=True)
class HouseStyle:
    """One agency's form: its labels, its punctuation, its date notation.

    **This is the variation the exercise is actually about.** A connector that
    reads only documents it generated in one house style proves nothing: the
    single hardest thing about a real inbox is that six agencies ask for the
    same twelve values under twelve different names, in four date notations,
    with and without colons, in one column or two.

    `labels` is deliberately *not* drawn from :data:`LABEL_SYNONYMS`. That table
    is shared with the extractor so the two cannot disagree about which labels
    exist — which makes any score measured on it circular. These houses include
    labels the extractor knows only from its own `EXTRA_SYNONYMS`, and one house
    (`cochin_pa`) whose labels nothing in the reader has ever been told about.
    Whatever that house loses is the honest residual, reported rather than
    tuned away.
    """

    key: str
    #: The agency or authority whose form this is, printed under the port's
    #: letterhead. `None` means the port authority's own form.
    filer: Optional[str]
    labels: dict
    #: Upper-case the labels, as a typewritten government form does.
    upper: bool = False
    #: Number the items ("1. Name of Ship"), as most Indian port forms do.
    numbered: bool = False
    #: What sits between a label and its value. `""` means the form used a
    #: column gap and no punctuation at all, which is what a PDF text layer
    #: hands back when the printed form had ruled boxes.
    separator: str = ":"
    #: Lay two label/value pairs on one line.
    two_column: bool = False
    #: Which notation :func:`eta_text` writes dates in; `None` draws one.
    date_style: Optional[int] = None


_HOUSE_LABELS: dict[str, dict] = {
    # 1. A port authority's own printed form. Every label here is one the
    #    reader already knows from the shared table — this house is the control.
    "deendayal_kandla": {
        "vessel_name": "Name of Vessel", "imo": "IMO Number",
        "call_sign": "Call Sign", "flag": "Flag",
        "last_port": "Last Port of Call", "arrival_port": "Port of Arrival",
        "eta": "ETA", "cargo": "Nature of Cargo", "crew_count": "No. of Crew",
        "owner": "Registered Owner", "agent": "Local Agent",
        "filed_at": "Date of Submission",
    },
    # 2. A terminal's form, written in the vocabulary of the IMO FAL forms.
    #    Every label is one the *extractor* knows and the *generator* has never
    #    written before, so this house tests the widening rather than rehearsing
    #    it.
    "mundra_terminal": {
        "vessel_name": "Name of Ship", "imo": "IMO Ship Identification Number",
        "call_sign": "Signal Letters", "flag": "Country of Registry",
        "last_port": "Previous Port of Call", "arrival_port": "Port of Destination",
        "eta": "Expected Date of Arrival", "cargo": "Description of Cargo",
        "crew_count": "Number of Crew", "owner": "Registered Owners",
        "agent": "Shipping Agent", "filed_at": "Date Filed",
    },
    # 3. A numbered statutory form, of the kind JNPA circulates.
    "jnpa_nhava_sheva": {
        "vessel_name": "Name of the Vessel", "imo": "IMO No.",
        "call_sign": "Int'l Call Sign", "flag": "Flag/Nationality",
        "last_port": "Port Sailed From", "arrival_port": "Port of Entry",
        "eta": "Arrival Date and Time", "cargo": "Cargo Particulars",
        "crew_count": "Persons on Board", "owner": "Owner/Operator",
        "agent": "Agent's Name", "filed_at": "Date & Time of Submission",
    },
    # 4. A telex-era form: no colons, columns held apart by spacing, every
    #    label as short as it can be. This is the house that exercises the
    #    reader's gap-pair path, which only accepts an exact label match.
    "nmpa_mangalore": {
        "vessel_name": "Vessel", "imo": "IMO", "call_sign": "C/S",
        "flag": "Flag State", "last_port": "From", "arrival_port": "To",
        "eta": "ETA", "cargo": "Commodity", "crew_count": "Crew",
        "owner": "Owners", "agent": "Agency", "filed_at": "Filed On",
    },
    # 5. **The wild house.** An older register's vocabulary: nothing in the
    #    reader has been told about "Whence Arrived" or "Souls on Board", and
    #    nothing will be for the sake of a number. What this house loses is the
    #    honest measure of how a genuinely unseen agency form reads.
    "cochin_pa": {
        "vessel_name": "Ship Identification", "imo": "Official Number",
        "call_sign": "Wireless Call Sign", "flag": "Register Nation",
        "last_port": "Whence Arrived", "arrival_port": "Port Applied For",
        "eta": "Probable Time of Arrival", "cargo": "Merchandise on Board",
        "crew_count": "Souls on Board", "owner": "Beneficial Owner",
        "agent": "Handling Agent", "filed_at": "Dated",
    },
    # 6. A shipping agent's own letterhead, typed in two columns at 0300.
    "agent_letterhead": {
        "vessel_name": "Vessel Name", "imo": "IMO No.", "call_sign": "Callsign",
        "flag": "Flag", "last_port": "Previous Port", "arrival_port": "Port",
        "eta": "Estimated Arrival", "cargo": "Cargo on Board",
        "crew_count": "Crew on Board", "owner": "Owners",
        "agent": "Agency", "filed_at": "Submitted",
    },
}

HOUSE_STYLES: dict[str, HouseStyle] = {
    "deendayal_kandla": HouseStyle(
        "deendayal_kandla", None, _HOUSE_LABELS["deendayal_kandla"],
        upper=True, numbered=True, separator=":", date_style=0),
    "mundra_terminal": HouseStyle(
        "mundra_terminal", "Adani Ports & SEZ — Marine Services",
        _HOUSE_LABELS["mundra_terminal"], separator=":", date_style=1),
    "jnpa_nhava_sheva": HouseStyle(
        "jnpa_nhava_sheva", None, _HOUSE_LABELS["jnpa_nhava_sheva"],
        numbered=True, separator=":", date_style=2),
    "nmpa_mangalore": HouseStyle(
        "nmpa_mangalore", None, _HOUSE_LABELS["nmpa_mangalore"],
        upper=True, separator="", date_style=1),
    "cochin_pa": HouseStyle(
        "cochin_pa", None, _HOUSE_LABELS["cochin_pa"],
        separator=":", date_style=3),
    "agent_letterhead": HouseStyle(
        "agent_letterhead", "Meridian Port Agents (I) Pvt Ltd",
        _HOUSE_LABELS["agent_letterhead"], separator=":", two_column=True,
        date_style=None),
}

#: Which house a port's paperwork is filed on, when the authoring does not say.
#: Ports without an entry cycle through the whole set.
_PORT_HOUSE = {
    "Kandla": "deendayal_kandla",
    "Mundra": "mundra_terminal",
    "JNPT": "jnpa_nhava_sheva",
    "Mumbai": "jnpa_nhava_sheva",
    "Mangalore": "nmpa_mangalore",
    "Kochi": "cochin_pa",
}


@dataclass
class NotificationSpec:
    """What one notification says. Values are strings, as a form holds them."""

    notification_id: str
    document_format: str
    received_at: datetime
    values: dict = field(default_factory=dict)
    #: The truth this document was written from, for the answer key. **Never
    #: written into the document and never read by any extractor** — same rule
    #: as `scenario_truth` (ADR-019).
    vessel_entity_id: Optional[str] = None
    #: Which fields were deliberately dropped, so the corpus can measure "the
    #: reader found what was there" separately from "the form was complete".
    omitted: tuple = ()
    #: Which kind of paperwork this is. Defaults to `pans` so every spec built
    #: before this existed renders exactly as it did — the PANS corpus ADR-036's
    #: figures were measured on must not move because a crew list was added.
    document_kind: str = "pans"
    #: Whose form it is. `None` keeps the original rendering: one house style,
    #: labels drawn from `LABEL_SYNONYMS`, no letterhead.
    house_style: Optional[str] = None
    #: The port whose authority prints the letterhead. `None` means no
    #: letterhead, which is the original behaviour.
    port: Optional[str] = None
    #: Rows the form carries below its header block — a crew table, manifest
    #: lines. Built by the authoring, not read by anything downstream.
    body_rows: tuple = ()
    #: Which case this document was authored as, for the answer key. **Never
    #: written into the document**, same rule as `vessel_entity_id`.
    authored_case: str = "honest"
    #: What the paperwork rules should say about it, per check. Answer key
    #: only; nothing in `ingest/` may read it.
    expected: dict = field(default_factory=dict)

    @property
    def document_name(self) -> str:
        ext = {"pdf": "pdf", "pdf_scan": "pdf", "docx": "docx",
               "xlsx": "xlsx", "electronic": "json"}[self.document_format]
        return f"{self.notification_id}.{ext}"

    @property
    def kind(self) -> DocumentKind:
        return DOCUMENT_KINDS[self.document_kind]


# --------------------------------------------------------------------------
# how a human types a vessel's name
# --------------------------------------------------------------------------

def mistype(name: str, rng: random.Random) -> tuple[str, str]:
    """A hand-typed rendering of a vessel name, and what kind of mangling it is.

    Four kinds, all of them things a clerk does at three in the morning, and
    each one breaking a *different* naive matcher:

    * a prefix — beats exact match, survives a normaliser that strips prefixes
    * a transposition — beats exact match and prefix stripping alike
    * a dropped space — beats anything tokenising on whitespace
    * nothing at all — the control, and the commonest case in real data

    The mix matters more than any single case: a corpus of only-mangled names
    would make an aggressive fuzzy matcher look good, and an aggressive fuzzy
    matcher is how a notification gets attached to the wrong hull.
    """
    kind = rng.choices(("clean", "prefix", "transpose", "nospace"),
                       weights=(0.55, 0.2, 0.15, 0.10))[0]
    if kind == "clean":
        return name, "clean"
    if kind == "prefix":
        return f"{rng.choice(('M.V. ', 'MV ', 'M/V '))}{name}", "prefix"
    if kind == "nospace":
        return name.replace(" ", ""), "nospace"
    # Transpose two adjacent letters inside a word — TRIUMPH -> TRUIMPH.
    letters = [i for i, c in enumerate(name[:-1])
               if c.isalpha() and name[i + 1].isalpha()]
    if not letters:
        return name, "clean"
    i = rng.choice(letters)
    return name[:i] + name[i + 1] + name[i] + name[i + 2:], "transpose"


# --------------------------------------------------------------------------
# writers, one per format
# --------------------------------------------------------------------------

_TITLE = "PRE-ARRIVAL NOTIFICATION OF SHIPS"
_SUBTITLE = "Form PANS-1  |  submitted under Merchant Shipping rules"


def _lines(spec: NotificationSpec, rng: random.Random) -> list[tuple[str, str]]:
    """(label, value) in this document's own field order, with its own labels.

    **The filing date is written into the document**, and it is written here
    rather than left to each scenario so that no notification can be produced
    without one. When it is absent, the only remaining evidence of when a form
    was filed is the file's modification time — which is when somebody scanned
    or copied it, not when the agent submitted it. Every rule that compares a
    declaration against the track needs the *filing* time, and a rule handed a
    scanning timestamp returns "not checkable" for a whole corpus while looking
    like it is working.
    """
    order = FIELD_ORDERS[rng.randrange(len(FIELD_ORDERS))]
    values = dict(spec.values)
    values.setdefault("filed_at", eta_text(spec.received_at, rng))
    out = []
    for key in order:
        if key in spec.omitted or key not in values:
            continue
        label = rng.choice(LABEL_SYNONYMS[key])
        out.append((label, str(values[key])))
    return out


# --------------------------------------------------------------------------
# a document as blocks — one layout, five renderings
# --------------------------------------------------------------------------

#: Wide enough that a PDF text layer keeps it as a column gap rather than
#: collapsing it to one space, which is what `extract._gap_pairs` needs to see.
_GAP = "     "


def _house_of(spec: NotificationSpec) -> Optional[HouseStyle]:
    return HOUSE_STYLES.get(spec.house_style) if spec.house_style else None


def _label_text(house: HouseStyle, field_name: str, index: int) -> str:
    label = house.labels.get(field_name) or field_name.replace("_", " ").title()
    if house.upper:
        label = label.upper()
    if house.numbered:
        label = f"{index}. {label}"
    return label


def _house_pairs(spec: NotificationSpec,
                 rng: random.Random) -> list[tuple[str, str]]:
    """(label, value) for one document, in its kind's order and house's words.

    A field the kind has no box for is not written, and a field the authoring
    dropped is not written. Those are different facts about a form — the first
    is what the form asks, the second is what the agent left out — and the
    answer key keeps them apart so "not checkable because a departure report has
    no ETA box" is never counted as "the reader missed it".
    """
    house = _house_of(spec)
    assert house is not None
    values = dict(spec.values)
    out: list[tuple[str, str]] = []
    index = 0
    for name in spec.kind.fields:
        if name in spec.omitted:
            continue
        value = values.get(name)
        if value in (None, ""):
            continue
        index += 1
        out.append((_label_text(house, name, index), str(value)))
    return out


def _head_lines(spec: NotificationSpec) -> list[str]:
    """The letterhead and title block, as an Indian port form prints it."""
    house = _house_of(spec)
    kind = spec.kind
    authority, address = authority_for(spec.port)
    lines = [authority, address]
    if house is not None and house.filer:
        lines.append(f"Filed by: {house.filer}")
    lines.append("")
    lines.append(kind.title)
    lines.append(f"{kind.form_ref}   |   Ref. No. {spec.notification_id}")
    if kind.note:
        lines.append(kind.note)
    return lines


def _pair_lines(spec: NotificationSpec,
                pairs: list[tuple[str, str]]) -> list[str]:
    """The header block as printed text — one column or two, colons or not."""
    house = _house_of(spec)
    assert house is not None
    sep = house.separator
    rendered = [f"{label}{sep} {value}" if sep else f"{label}{_GAP}{value}"
                for label, value in pairs]
    if not house.two_column:
        return rendered
    out = []
    for i in range(0, len(rendered), 2):
        out.append(_GAP.join(rendered[i:i + 2]))
    return out


def _body_table(spec: NotificationSpec) -> Optional[tuple[tuple, tuple]]:
    """(headers, rows) for the table this kind of form carries, or None."""
    if not spec.body_rows:
        return None
    headers = _BODY_HEADERS.get(spec.kind.body)
    if headers is None:
        return None
    return headers, tuple(spec.body_rows)


#: Column headings for each body. Taken from the IMO FAL forms these documents
#: are modelled on, including "Nationality" — which is a label the reader also
#: knows as a synonym for the vessel's flag. That collision is left in on
#: purpose: it is in the real form, and a corpus that quietly renamed it would
#: hide a misattribution the reader has to be able to survive.
_BODY_HEADERS: dict[str, tuple] = {
    "crew": ("No.", "Family Name", "Given Names", "Rank", "Nationality",
             "Date of Birth", "Passport No."),
    "manifest": ("Item", "Marks & Nos.", "No. of Packages",
                 "Description of Goods", "Gross Weight (MT)"),
}

_CLEARANCE_TEXT = (
    "Clearance is hereby granted to the above named vessel to proceed to sea, "
    "the dues and charges leviable under the Merchant Shipping Act, 1958 "
    "having been paid and all required documents having been produced.")


def _footer_lines(spec: NotificationSpec) -> list[str]:
    """The sign-off block every one of these forms carries.

    It is written because real forms have one and because it is a trap worth
    keeping: it repeats the agent's name under a label nothing recognises, and
    it holds the word "Master", which the extractor recognises *in order to
    refuse it*. A form whose footer changed the agent field would be a
    misattribution, and the corpus has to be able to catch that.
    """
    agent = spec.values.get("agent") or "the Agent"
    lines = ["", "I declare that the particulars given above are true and "
                 "correct to the best of my knowledge and belief.",
             f"For and on behalf of: {agent}",
             "Signature ____________________   Master / Authorised Agent"]
    if spec.kind.body == "clearance":
        lines = ["", _CLEARANCE_TEXT] + lines[1:]
    return lines


def _document_lines(spec: NotificationSpec,
                    rng: random.Random) -> list[str]:
    """The whole document as printed lines — what a PDF or a fax shows."""
    pairs = _house_pairs(spec, rng)
    lines = _head_lines(spec) + [""] + _pair_lines(spec, pairs)
    table = _body_table(spec)
    if table is not None:
        headers, rows = table
        lines.append("")
        lines.append(_GAP.join(str(h) for h in headers))
        for row in rows:
            lines.append(_GAP.join(str(c) for c in row))
    lines.extend(_footer_lines(spec))
    return lines


def _write_pdf(path: Path, spec: NotificationSpec, rng: random.Random) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    if spec.house_style is not None:
        return _write_pdf_styled(path, spec, rng)

    c = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    y = height - 70
    c.setFont("Helvetica-Bold", 13)
    c.drawString(60, y, _TITLE)
    y -= 18
    c.setFont("Helvetica", 8)
    c.drawString(60, y, _SUBTITLE)
    y -= 30
    c.setFont("Helvetica", 10)
    for label, value in _lines(spec, rng):
        c.drawString(60, y, f"{label}: {value}")
        y -= 17
    c.save()


def _write_pdf_styled(path: Path, spec: NotificationSpec,
                      rng: random.Random) -> None:
    """A letterheaded form, in a monospaced face.

    Monospaced because the telex-era house style separates its columns with
    spacing and nothing else, and a proportional face renders that gap narrow
    enough that a PDF text layer hands back one space — at which point the
    reader sees a sentence rather than a label and a value. The typeface is a
    property of the *test*, not a hidden helper: real forms of that vintage are
    monospaced for exactly the same reason.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 60
    for i, line in enumerate(_document_lines(spec, rng)):
        if y < 50:
            c.showPage()
            y = height - 60
        c.setFont("Courier-Bold" if i in (0, 4) else "Courier",
                  9 if len(line) > 96 else 10)
        c.drawString(45, y, line)
        y -= 14
    c.save()


def _render_page_image(spec: NotificationSpec, rng: random.Random):
    """A page rendered as pixels — the thing a scanner produces.

    Deliberately imperfect: a slight rotation and a little noise, because a
    scanned page is never square on the glass and a clean render would be an
    OCR problem this project does not have. The tilt is small enough that
    tesseract copes, which is the point — the corpus has to be hard enough to
    be real and not so hard that it measures tesseract instead of the pipeline.
    """
    from PIL import Image, ImageDraw, ImageFont

    if spec.house_style is None:
        W, H, size, step = 1400, 1000, 26, 46
        rendered = None
    else:
        # A letterheaded form is a full page of text — a crew list runs to
        # thirty lines — so it is rendered at something near A4 at 200 dpi. The
        # old 1400x1000 sheet fitted twelve lines and would have silently
        # cropped the crew off the bottom of every crew list, which reads
        # downstream as a form the agent left half empty.
        W, H, size, step = 1700, 2340, 23, 36
        rendered = _document_lines(spec, rng)
    img = Image.new("L", (W, H), 255)
    d = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 30)
        font = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"
            if rendered is not None else
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", size)
    except OSError:                                          # pragma: no cover
        title_font = font = ImageFont.load_default()

    if rendered is None:
        d.text((60, 50), _TITLE, fill=20, font=title_font)
        y = 130
        rendered = [f"{label}: {value}" for label, value in _lines(spec, rng)]
    else:
        y = 50
    for line in rendered:
        d.text((60, y), line, fill=30, font=font)
        y += step
        if y > H - 60:
            break

    img = img.rotate(rng.uniform(-0.7, 0.7), fillcolor=255, resample=Image.BICUBIC)
    return img


def _write_pdf_scan(path: Path, spec: NotificationSpec,
                    rng: random.Random) -> None:
    """A PDF with **no text layer** — an image of a page, as a fax produces.

    This is the format the requirement singles out and the one that decides
    whether the extractor is real. `pypdf` returns an empty string for it, so a
    reader that only knows how to pull text sees a blank document and reports a
    notification with no fields rather than an unreadable one — which is a much
    worse failure than an error, because it looks like a form somebody left
    empty.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    img = _render_page_image(spec, rng)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    w, h = A4
    c = canvas.Canvas(str(path), pagesize=A4)
    c.drawImage(ImageReader(buf), 0, 0, width=w, height=h, preserveAspectRatio=True)
    c.save()


def _write_docx(path: Path, spec: NotificationSpec, rng: random.Random) -> None:
    """Word, and **half of them as a table** rather than as paragraphs.

    Both shapes are common and they need different reading. A reader that walks
    paragraphs finds nothing in a table document, and the failure is silent for
    the same reason the scan's is: it produces an empty form, not an error.
    """
    from docx import Document

    if spec.house_style is not None:
        return _write_docx_styled(path, spec, rng)

    doc = Document()
    doc.add_heading(_TITLE, level=1)
    doc.add_paragraph(_SUBTITLE)
    rows = _lines(spec, rng)
    if rng.random() < 0.5:
        table = doc.add_table(rows=0, cols=2)
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
    else:
        for label, value in rows:
            doc.add_paragraph(f"{label}: {value}")
    doc.save(str(path))


def _write_docx_styled(path: Path, spec: NotificationSpec,
                       rng: random.Random) -> None:
    """A letterheaded Word form — header block as paragraphs or as a table.

    The body (a crew list, a manifest) is always a real Word table, because
    that is how it arrives: nobody types thirty crew as paragraphs. That makes
    the body the place where a table-shaped reader can go wrong, which is what
    it is here to test.
    """
    from docx import Document

    doc = Document()
    head = _head_lines(spec)
    doc.add_heading(spec.kind.title, level=1)
    for line in head:
        if line and line != spec.kind.title:
            doc.add_paragraph(line)

    pairs = _house_pairs(spec, rng)
    house = _house_of(spec)
    if house is not None and (house.two_column or rng.random() < 0.5):
        table = doc.add_table(rows=0, cols=2)
        for label, value in pairs:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
    else:
        for line in _pair_lines(spec, pairs):
            doc.add_paragraph(line)

    body = _body_table(spec)
    if body is not None:
        headers, rows = body
        table = doc.add_table(rows=0, cols=len(headers))
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(h)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
    for line in _footer_lines(spec):
        if line:
            doc.add_paragraph(line)
    doc.save(str(path))


def _write_xlsx(path: Path, spec: NotificationSpec, rng: random.Random) -> None:
    """A spreadsheet, with the form starting a few rows down.

    Real submitted spreadsheets have a logo, a blank row and a heading before
    the data starts, so the label column is not column A and the first row is
    not the header. A reader that assumes A1 reads the logo.
    """
    from openpyxl import Workbook

    if spec.house_style is not None:
        return _write_xlsx_styled(path, spec, rng)

    wb = Workbook()
    ws = wb.active
    ws.title = "PANS"
    ws["B2"] = _TITLE
    ws["B3"] = _SUBTITLE
    r = 5 + rng.randrange(0, 3)
    col_label = rng.choice(("A", "B", "C"))
    col_value = chr(ord(col_label) + 1)
    for label, value in _lines(spec, rng):
        ws[f"{col_label}{r}"] = label
        ws[f"{col_value}{r}"] = value
        r += 1
    wb.save(str(path))


def _write_xlsx_styled(path: Path, spec: NotificationSpec,
                       rng: random.Random) -> None:
    """The letterheaded form as a submitted workbook.

    The body goes on **a second sheet**, which is what an agent's template
    does — and it matters to the reader, because `read_xlsx` walks every
    worksheet and pairs adjacent cells on every row. A crew table read that way
    produces passages like "Rank: Nationality": a column heading beside a
    column heading. Nothing downstream may read one of those as an answer, and
    this is the corpus that proves it does not.
    """
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = spec.kind.prefix
    r = 1
    for line in _head_lines(spec):
        if line:
            ws.cell(row=r, column=2, value=line)
        r += 1
    r += 1 + rng.randrange(0, 2)
    col_label = rng.randrange(1, 4)
    for label, value in _house_pairs(spec, rng):
        ws.cell(row=r, column=col_label, value=label)
        ws.cell(row=r, column=col_label + 1, value=value)
        r += 1

    body = _body_table(spec)
    if body is not None:
        headers, rows = body
        ws2 = wb.create_sheet(title={"crew": "Crew", "manifest": "Manifest"}
                              .get(spec.kind.body, "Detail"))
        for i, h in enumerate(headers, start=1):
            ws2.cell(row=1, column=i, value=str(h))
        for j, row in enumerate(rows, start=2):
            for i, value in enumerate(row, start=1):
                ws2.cell(row=j, column=i, value=str(value))
    r += 1
    for line in _footer_lines(spec):
        if line:
            ws.cell(row=r, column=1, value=line)
        r += 1
    wb.save(str(path))


def _write_electronic(path: Path, spec: NotificationSpec,
                      _rng: random.Random) -> None:
    """The portal feed: structured, complete, and boring.

    *"Structure ingestion so the electronic feed drops in as an additional
    source without rework, since the requirement asks for that compatibility
    explicitly."* It is generated here so that claim is tested by a reader
    producing the same record as the document readers, rather than argued for
    in a design note.

    Note it keeps the **same hand-typed vessel name** as the documents. A
    structured feed removes format ambiguity; it does not stop an agent typing
    "GRANITE TRUIMPH" into a web form.
    """
    def declared(name):
        """The value the portal would carry, honouring the kind and the drops.

        A portal payload for a departure report has no arrival-port key,
        because the form it is a rendering of has no arrival-port box. That
        keeps "not checkable" meaning the same thing on the electronic path as
        it does on paper — which is the whole of the compatibility claim.
        """
        if name in spec.omitted or name not in spec.kind.fields:
            return None
        return spec.values.get(name)

    payload = {
        "notificationId": spec.notification_id,
        # **The portal says what kind of document it is sending.** A real
        # national-logistics feed carries a document type, and without one the
        # electronic path is the only path where the kind has to be guessed
        # from a filename. It is a passage like any other, not a field.
        "documentType": spec.kind.title,
        "formReference": spec.kind.form_ref,
        "submittedAt": spec.received_at.isoformat(),
        "vessel": {
            "name": declared("vessel_name"),
            "imo": declared("imo"),
            "callSign": declared("call_sign"),
            "flag": declared("flag"),
        },
        "voyage": {
            "lastPort": declared("last_port"),
            "arrivalPort": declared("arrival_port"),
            "eta": declared("eta"),
            "cargo": declared("cargo"),
        },
        "parties": {
            "owner": declared("owner"),
            "agent": declared("agent"),
            "crew": declared("crew_count"),
        },
    }
    # Omitted fields are absent, not null — a portal that sends `null` and one
    # that omits the key are saying the same thing, and the reader has to cope
    # with the shape that is harder to notice.
    for group in ("vessel", "voyage", "parties"):
        payload[group] = {k: v for k, v in payload[group].items()
                          if v is not None}
    path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


_WRITERS = {
    "pdf": _write_pdf,
    "pdf_scan": _write_pdf_scan,
    "docx": _write_docx,
    "xlsx": _write_xlsx,
    "electronic": _write_electronic,
}


def write_notifications(specs, out_dir: Path, *, seed: int = 7) -> dict:
    """Write every spec to disk in its own format. Returns a per-format count.

    Missing writers are reported rather than raised: a laptop without
    `python-docx` should still produce the other four formats and say which one
    it could not write, because a generator that refuses to run is a generator
    nobody runs.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    written: dict[str, int] = {}
    skipped: dict[str, str] = {}
    for spec in specs:
        writer = _WRITERS[spec.document_format]
        rng = random.Random(
            int(hashlib.sha256(spec.notification_id.encode()).hexdigest()[:8], 16)
            ^ seed)
        try:
            writer(out_dir / spec.document_name, spec, rng)
        except ImportError as e:                             # pragma: no cover
            skipped.setdefault(spec.document_format, str(e))
            continue
        written[spec.document_format] = written.get(spec.document_format, 0) + 1
    if skipped:
        written["_skipped"] = skipped                        # pragma: no cover
    return written


def eta_text(when: datetime, rng: random.Random) -> str:
    """An ETA written the way a form holds it, which is never ISO 8601.

    Four renderings, all real: day-first with slashes, day-first with a month
    name, ISO date with a separate time, and a 12-hour clock. A reader that
    parses one of these parses a quarter of an inbox.
    """
    style = rng.randrange(4)
    if style == 0:
        return when.strftime("%d/%m/%Y %H:%M")
    if style == 1:
        return when.strftime("%d %b %Y %H%M") + " LT"
    if style == 2:
        return when.strftime("%Y-%m-%d") + " at " + when.strftime("%H:%M")
    hour = when.strftime("%I:%M %p").lstrip("0")
    return when.strftime("%d-%m-%Y") + " " + hour


def jitter_eta(when: datetime, rng: random.Random,
               *, hours: float = 6.0) -> datetime:
    """A declared ETA near the real arrival, because paperwork is filed early.

    A notification is submitted 24-72 hours out and states an estimate. Making
    the declared time exactly equal the observed arrival would build a corpus
    where any disagreement at all is a finding, and the arrival-window rule
    would then be measuring the generator's precision rather than a vessel's
    honesty.
    """
    return when + timedelta(hours=rng.uniform(-hours, hours))


def _date_style(when: datetime, style: int) -> str:
    """The four notations of :func:`eta_text`, chosen rather than drawn.

    A house style is a *habit*: an agency writes its dates the same way every
    time, and drawing one per document would model an agency that cannot decide.
    The variety across the corpus comes from there being six houses, which is
    where it comes from in a real inbox.
    """
    style = style % 4
    if style == 0:
        return when.strftime("%d/%m/%Y %H:%M")
    if style == 1:
        return when.strftime("%d %b %Y %H%M") + " LT"
    if style == 2:
        return when.strftime("%Y-%m-%d") + " at " + when.strftime("%H:%M")
    hour = when.strftime("%I:%M %p").lstrip("0")
    return when.strftime("%d-%m-%Y") + " " + hour


# ==========================================================================
# the wider corpus: six kinds of paper, six houses, and the cases they carry
# ==========================================================================

#: What a document was authored to be. **Never written into a document** — the
#: same rule `vessel_entity_id` and `scenario_truth` live under (ADR-019). It
#: exists so a run can report "of the eleven forms authored to contradict a
#: track, the rules found nine", which is the only way an extraction figure
#: means anything.
AUTHORED_CASES: tuple[str, ...] = (
    "honest",
    "false_last_port",
    "missed_arrival_window",
    "false_ballast",
    "absent_last_port",
    "absent_eta",
    "declared_absence_cargo",
    "unresolvable_hull",
)

#: One cycle of twenty documents. Mostly honest, because **the honest majority
#: is the denominator**: a paperwork rule measured only against forged forms
#: reports recall and says nothing about how often it accuses an agent who
#: typed a date wrong. The cycle is walked in order rather than drawn, so a
#: corpus of forty documents contains exactly two of every authored case and a
#: per-case figure is never one document wide.
CASE_MIX: tuple[str, ...] = (
    "honest", "honest", "false_last_port", "honest", "absent_eta",
    "honest", "missed_arrival_window", "honest", "declared_absence_cargo",
    "honest", "false_ballast", "honest", "absent_last_port", "honest",
    "missed_arrival_window", "honest", "unresolvable_hull", "honest",
    "false_last_port", "honest",
)

#: Which authored cases a kind of form can carry. A departure report has no
#: ETA box, so it cannot be authored to miss an arrival window — and a rule
#: that answered anything but "not checkable" on one would be wrong.
_CASE_NEEDS = {
    "false_last_port": ("last_port",),
    "absent_last_port": ("last_port",),
    "missed_arrival_window": ("eta", "arrival_port"),
    "absent_eta": ("eta",),
    "false_ballast": ("cargo",),
    "declared_absence_cargo": ("cargo",),
}

#: Cargo descriptions a form carries, as free text.
CARGOES: tuple[str, ...] = (
    "Crude oil in bulk", "Gas oil", "Containerised general cargo",
    "Bulk cement", "Iron ore fines", "Refined products", "Bulk fertiliser",
    "Project cargo", "Frozen fish", "Bunkers only", "Rock phosphate",
    "Edible oil in bulk", "Steel coils", "Bagged sugar",
)

#: The one declaration `anomaly.paperwork.check_declared_ballast` can check.
BALLAST_DECLARATION = "Ballast — no cargo"

#: Agencies that file on a vessel's behalf.
AGENCIES: tuple[str, ...] = (
    "Coastal Marine Services", "Anchor Shipping Agency",
    "Meridian Port Agents (I) Pvt Ltd", "Gateway Marine Pvt Ltd",
    "Sagar Shipping Services", "Konkan Maritime Agencies",
    "Saurashtra Ship Agency", "Malabar Marine Services",
)

_SURNAMES = ("Nair", "Fernandes", "Reyes", "Bautista", "Kovalenko", "Singh",
             "D'Souza", "Pillai", "Cruz", "Shevchenko", "Rathore", "Menon",
             "Villanueva", "Bose", "Iqbal", "Thapa", "Gomes", "Kulkarni")
_GIVEN = ("Rajesh", "Anil", "Miguel", "Ramon", "Oleksandr", "Harpreet",
          "Vinod", "Suresh", "Jose", "Dmytro", "Arjun", "Sandeep",
          "Ferdinand", "Prakash", "Imran", "Bikash", "Antonio", "Girish")
_RANKS = ("Master", "Chief Officer", "Second Officer", "Third Officer",
          "Chief Engineer", "Second Engineer", "Third Engineer", "Electrician",
          "Bosun", "Able Seaman", "Able Seaman", "Ordinary Seaman", "Oiler",
          "Fitter", "Cook", "Messman", "Deck Cadet", "Engine Cadet")
_NATIONALITIES = ("Indian", "Filipino", "Ukrainian", "Indonesian",
                  "Bangladeshi", "Sri Lankan", "Myanmarese")


def _crew_rows(count: int, rng: random.Random) -> tuple:
    """A crew table, one line per soul on board, as IMO FAL Form 5 holds it."""
    rows = []
    for i in range(1, max(1, count) + 1):
        rank = _RANKS[(i - 1) % len(_RANKS)]
        rows.append((
            str(i), rng.choice(_SURNAMES), rng.choice(_GIVEN), rank,
            rng.choice(_NATIONALITIES),
            f"{rng.randint(1, 28):02d}/{rng.randint(1, 12):02d}/"
            f"{rng.randint(1965, 2003)}",
            f"{rng.choice('ABGHJKMNPRSTZ')}{rng.randint(1000000, 9999999)}",
        ))
    return tuple(rows)


def _manifest_rows(cargo: str, rng: random.Random) -> tuple:
    """A handful of manifest lines under the declared commodity."""
    rows = []
    for i in range(1, rng.randint(3, 6)):
        rows.append((
            str(i), f"{rng.choice('ABCDEFGH')}{rng.randint(100, 999)}/"
                    f"{rng.randint(1, 40)}",
            str(rng.randint(20, 1800)),
            cargo if i == 1 else rng.choice(CARGOES),
            f"{rng.randint(150, 42000):,}",
        ))
    return tuple(rows)


def _far_port(port: Optional[str]) -> Optional[str]:
    """The gazetteer port furthest from this one — a last port she was not at.

    Computed rather than listed, so a corpus whose arrivals move does not need
    a hand-maintained table of lies. The furthest port in this gazetteer is
    always more than a thousand kilometres away, comfortably outside
    `paperwork.LAST_PORT_RADIUS_KM`, so the contradiction is unambiguous rather
    than marginal — a marginal one would be measuring the radius rather than
    the rule.
    """
    from ..ports import PORTS

    if not port or port not in PORTS:
        return None
    plat, plon = PORTS[port]
    best, best_d = None, -1.0
    for name, (lat, lon) in sorted(PORTS.items()):
        d = (lat - plat) ** 2 + (lon - plon) ** 2
        if d > best_d:
            best, best_d = name, d
    return best


def _kind_supports(kind: DocumentKind, case: str) -> bool:
    return all(f in kind.fields for f in _CASE_NEEDS.get(case, ()))


def _take_case(pending: list, kind: DocumentKind) -> str:
    """The next case from the mix that this kind of form can actually carry.

    **A case a kind cannot carry is put back, never consumed.** This looks like
    a detail and is not: the earlier version walked forward through
    :data:`CASE_MIX` and *spent* every case it skipped. Six kinds cycle against
    a twenty-slot mix, so the two `missed_arrival_window` slots landed on a
    port clearance and a departure report — neither of which has an ETA box —
    on every single lap. The result was a 161-document corpus in which that
    case appeared **zero** times, and therefore a corpus that never once
    exercised the arrival-window contradiction the requirement names
    explicitly.

    It failed silently, which is the failure mode this whole area is built
    against: nothing errored, every document was valid, the read rates looked
    fine, and the run reported a per-case table with a row simply missing from
    it. Deferring instead of consuming means a case waits for the next kind of
    form that has the box for it, so the mix is delivered rather than merely
    intended.

    `pending` is mutated: it is the queue carried across documents, refilled a
    lap of the mix at a time.
    """
    for _ in range(2):
        for i, case in enumerate(pending):
            if _kind_supports(kind, case):
                return pending.pop(i)
        pending.extend(CASE_MIX)
    return "honest"                                          # pragma: no cover


def build_document_specs(voyages, *, seed: int = 11,
                         taken_imos=(), start_index: int = 1) -> list:
    """One document per voyage, over **whatever vessels the corpus holds**.

    `voyages` is a list of dicts, each describing one recorded arrival:

    ``vessel_id``, ``name``, ``imo``, ``call_sign``, ``flag``, ``owner``,
    ``arrival_port`` (a gazetteer name), ``arrival_time``, ``last_port`` (the
    port she actually sailed from, or ``None``), ``prior_call_end`` and
    ``draught_m``.

    **Nothing here knows any vessel's name in advance.** The cast is minted by
    another part of the generator and grows; a document set that named hulls
    would break the day the fleet changed, and would be describing a corpus
    rather than reading one.

    Kinds, formats and houses are cycled rather than drawn, so every reader,
    every form and every agency gets a comparable sample at any corpus size —
    the same reasoning `group_p.build_notifications` gives for round-robin
    formats. The *case* is cycled too, through :data:`CASE_MIX`, and skipped
    forward when the kind in hand has no box for the field the case needs.
    """
    rng = random.Random(seed ^ 0xD0C5)
    taken = {str(i) for i in taken_imos if i}
    specs: list[NotificationSpec] = []
    n = start_index
    houses = sorted(HOUSE_STYLES)
    #: Cases waiting to be authored. A case a kind cannot carry is **put back**,
    #: not consumed — see `_take_case`.
    pending: list[str] = []

    for i, v in enumerate(sorted(voyages,
                                 key=lambda d: (str(d.get("arrival_time")),
                                                str(d.get("vessel_id"))))):
        kind = DOCUMENT_KINDS[DOCUMENT_KIND_ORDER[i % len(DOCUMENT_KIND_ORDER)]]
        fmt = FORMATS[i % len(FORMATS)]
        port = v.get("arrival_port")
        house = _PORT_HOUSE.get(str(port)) or houses[i % len(houses)]

        case = _take_case(pending, kind)
        if case == "false_ballast" and not _is_laden(v):
            # She is not broadcasting a laden draught, so there is nothing for
            # the ballast rule to contradict. The case goes back in the queue
            # for the next hull that is laden, rather than being spent on one
            # where it could only ever read "not checkable".
            pending.insert(0, case)
            case = "honest"

        spec = _author(v, kind=kind, fmt=fmt, house=house, case=case,
                       index=n, rng=rng, taken=taken)
        if spec is None:
            continue
        specs.append(spec)
        n += 1
    return specs


def _is_laden(v) -> bool:
    """Is she broadcasting a draught only a loaded hull draws?

    The threshold is stated here rather than imported from `anomaly.paperwork`,
    for the reason `group_p` gives for the same duplication: a corpus built from
    a rule's own constant cannot falsify that rule. `tests/` asserts the two
    have not drifted, which is what makes independence safe rather than merely
    separate.
    """
    d = v.get("draught_m")
    try:
        return d is not None and float(d) >= 12.0
    except (TypeError, ValueError):                          # pragma: no cover
        return False


#: How far ahead of arrival a form is filed. The statutory window is 24-96
#: hours; the spread matters because the arrival-window rule measures the gap
#: between declaration and observation, and a fixed lead would make it constant.
FILE_LEAD_HOURS = (24.0, 96.0)

#: How far a form authored to miss its window misses it by. Well past
#: `paperwork.ARRIVAL_SLIP_HOURS` (24), because a case sitting on the threshold
#: measures the threshold rather than the rule.
MISSED_WINDOW_HOURS = (72.0, 160.0)


def _author(v, *, kind: DocumentKind, fmt: str, house: str, case: str,
            index: int, rng: random.Random,
            taken: set) -> Optional[NotificationSpec]:
    """One document: what it says, when it was filed, and what it is for."""
    arrival = v.get("arrival_time")
    if arrival is None:
        return None
    style = HOUSE_STYLES[house].date_style
    lead = rng.uniform(*FILE_LEAD_HOURS)
    filed = arrival - timedelta(hours=lead)
    prior_end = v.get("prior_call_end")
    if prior_end is not None and filed < prior_end:
        # She cannot declare a last port she has not reached yet. Filing starts
        # once she is alongside — the same correction `group_p` had to make
        # after authoring forms whose stated origin was, at filing time, in the
        # vessel's future.
        filed = prior_end + timedelta(hours=rng.uniform(0.5, 4.0))
        if filed >= arrival:
            return None

    typed, _mangle = mistype(str(v.get("name") or ""), rng)
    declared_eta = jitter_eta(arrival, rng, hours=5.0)
    last_port = v.get("last_port") or "Sea"
    cargo = rng.choice(CARGOES)
    if _is_laden(v) is False and rng.random() < 0.18:
        cargo = BALLAST_DECLARATION
    omitted: list[str] = []
    expected: dict = {}
    entity = v.get("vessel_id")

    if case == "false_last_port":
        far = _far_port(v.get("arrival_port"))
        if far is None:
            case = "honest"
        else:
            last_port = far
            expected["declared_last_port"] = "contradiction"
    if case == "missed_arrival_window":
        slip = rng.uniform(*MISSED_WINDOW_HOURS) * rng.choice((-1.0, 1.0))
        declared_eta = arrival + timedelta(hours=slip)
        expected["declared_arrival_window"] = "contradiction"
    if case == "false_ballast":
        cargo = BALLAST_DECLARATION
        expected["declared_ballast"] = "contradiction"
    if case == "absent_last_port":
        omitted.append("last_port")
        expected["declared_last_port"] = "not_checkable"
    if case == "absent_eta":
        omitted.append("eta")
        expected["declared_arrival_window"] = "not_checkable"
    if case == "declared_absence_cargo":
        # **"NIL" is not "in ballast".** ADR-036 is explicit that an agent's
        # dash means "I did not answer", and reading it as a cargo declaration
        # would fire the ballast rule on a hull nobody accused. The expected
        # outcome is therefore `not_checkable`, and a corpus run that reports a
        # contradiction here has found a real defect.
        cargo = rng.choice(("NIL", "N/A", "--", "TBA"))
        expected["declared_ballast"] = "not_checkable"

    if case == "unresolvable_hull":
        typed = _invented_name(rng)
        entity = None
        values_identity = dict(
            vessel_name=typed, imo=_free_imo(rng, taken),
            call_sign=_invented_call_sign(rng), flag=rng.choice(
                ("TGO", "COM", "TZA", "MNG", "CMR")))
        owner = f"{typed.title().split()[0]} Lines FZE"
    else:
        values_identity = dict(
            vessel_name=typed, imo=str(v.get("imo") or ""),
            call_sign=str(v.get("call_sign") or ""),
            flag=str(v.get("flag") or ""))
        owner = v.get("owner") or f"{str(v.get('name') or '').title()} Shipping Ltd"

    values = dict(values_identity)
    values.update(
        last_port=last_port,
        arrival_port=v.get("arrival_port") or "",
        eta=_date_style(declared_eta, style if style is not None
                        else rng.randrange(4)),
        cargo=cargo,
        crew_count=str(rng.randint(12, 28)),
        owner=owner,
        agent=rng.choice(AGENCIES),
        filed_at=_date_style(filed, style if style is not None
                             else rng.randrange(4)),
    )
    values = {k: val for k, val in values.items() if val not in (None, "")}

    # Real forms arrive incomplete for reasons nobody authored. A field dropped
    # here is *not* an authored case — it is the background incompleteness that
    # stops the reader being measured against a corpus where every box is
    # always filled — so it never writes an expectation.
    for name in ("owner", "call_sign", "crew_count", "agent"):
        if name not in omitted and rng.random() < 0.12:
            omitted.append(name)

    body: tuple = ()
    if kind.body == "crew":
        try:
            n_crew = int(values.get("crew_count") or 18)
        except ValueError:                                   # pragma: no cover
            n_crew = 18
        body = _crew_rows(min(n_crew, 24), rng)
    elif kind.body == "manifest":
        body = _manifest_rows(values.get("cargo") or "General cargo", rng)

    return NotificationSpec(
        notification_id=f"{kind.prefix}-{index:04d}",
        document_format=fmt,
        received_at=filed,
        values=values,
        vessel_entity_id=entity,
        omitted=tuple(dict.fromkeys(omitted)),
        document_kind=kind.key,
        house_style=house,
        port=v.get("arrival_port"),
        body_rows=body,
        authored_case=case,
        expected=expected,
    )


_INVENTED_FIRST = ("SEA", "OCEAN", "STAR", "SILVER", "GOLDEN", "NORTHERN",
                   "EASTERN", "CRYSTAL", "PACIFIC", "ORIENT")
_INVENTED_SECOND = ("HARRIER", "SPIRIT", "VENTURE", "TRADER", "PIONEER",
                    "SENTINEL", "MARINER", "FORTUNE", "HORIZON", "LEGACY")
_INVENTED_THIRD = ("PRIDE", "II", "EXPRESS", "GLORY", "STAR", "")


def _invented_name(rng: random.Random) -> str:
    """A hull nothing in the picture holds — the other half of the gap.

    *"A notification that cannot be matched to any track, or a vessel arriving
    with no notification at all, are both exactly the kind of gap the Coast
    Guard wants surfaced."* The name is built from ordinary merchant-ship
    words on purpose: an obviously fake name would be resolved by a human
    before it reached the resolver, and the finding is about a well-formed
    document naming a ship we cannot identify.
    """
    parts = [rng.choice(_INVENTED_FIRST), rng.choice(_INVENTED_SECOND),
             rng.choice(_INVENTED_THIRD)]
    return " ".join(p for p in parts if p)


def _invented_call_sign(rng: random.Random) -> str:
    return (rng.choice("9AVD3") + rng.choice("XQWVZ")
            + rng.choice("ABCDEFGH") + str(rng.randint(1, 9)))


def _free_imo(rng: random.Random, taken: set) -> str:
    """Seven digits nothing in the corpus answers to.

    Checked against the identifiers the caller passed in, because an "invented"
    IMO that happens to belong to a real hull would resolve — and the document
    authored to be unmatchable would land on a ship, quietly, as a paperwork
    finding against her.
    """
    for _ in range(500):
        candidate = str(rng.randint(1900000, 1999999))
        if candidate not in taken:
            taken.add(candidate)
            return candidate
    return str(rng.randint(1900000, 1999999))                # pragma: no cover
