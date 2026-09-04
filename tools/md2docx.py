"""Minimal markdown -> docx for the iDEX annexures.

Handles: ATX headings, paragraphs, bullet lists, ordered lists, pipe tables,
fenced code blocks, horizontal rules, and inline **bold** / *italic* / `code`.
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)", re.S)


def add_runs(par, text):
    text = text.replace("\n", " ")
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2])
            r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            r = par.add_run(piece[1:-1])
            r.italic = True
        else:
            par.add_run(piece)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def build(md_path, docx_path, title):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(10)
            par.paragraph_format.space_before = Pt(6)
            r = par.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(7.5)
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and re.match(
                r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for row in rows:
                cells = table.add_row().cells
                for c in range(len(header)):
                    text = row[c] if c < len(row) else ""
                    par = cells[c].paragraphs[0]
                    par.text = ""
                    par.paragraph_format.space_after = Pt(2)
                    add_runs(par, text)
                    for run in par.runs:
                        run.font.size = Pt(9)
            doc.add_paragraph()
            continue

        if not stripped:
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(2)
            r = par.add_run("_" * 96)
            r.font.size = Pt(6)
            r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                par = doc.add_paragraph()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(par, text)
                for run in par.runs:
                    run.bold = True
                    run.font.size = Pt(16)
            else:
                par = doc.add_heading(level=min(level, 4))
                par.text = ""
                add_runs(par, text)
                size = {2: 13, 3: 11.5, 4: 10.5}.get(level, 10.5)
                for run in par.runs:
                    run.font.size = Pt(size)
                    run.font.color.rgb = RGBColor(0x1F, 0x30, 0x51)
                    run.font.name = "Calibri"
            i += 1
            continue

        m = re.match(r"^([-*])\s+(.*)$", stripped) or re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            bullet = not m.group(1).isdigit()
            indent = len(line) - len(line.lstrip(" "))
            buf = [m.group(2)]
            i += 1
            # continuation lines: indented further than the marker, not a new marker
            while i < n:
                nxt = lines[i]
                nxt_s = nxt.strip()
                if not nxt_s or nxt_s.startswith("#") or nxt_s.startswith("|") \
                        or nxt_s.startswith("```") or nxt_s in ("---", "***", "___") \
                        or re.match(r"^([-*])\s+", nxt_s) or re.match(r"^\d+\.\s+", nxt_s):
                    break
                buf.append(nxt_s)
                i += 1
            style = "List Bullet" if bullet else "List Number"
            if indent >= 2:
                style = ("List Bullet 2" if bullet else "List Number 2")
            try:
                par = doc.add_paragraph(style=style)
            except KeyError:
                par = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            par.paragraph_format.space_after = Pt(3)
            add_runs(par, " ".join(buf))
            continue

        # paragraph: gather continuation lines
        buf = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("|")
                    or nxt.startswith("```") or nxt in ("---", "***", "___")
                    or re.match(r"^([-*])\s+", nxt)
                    or re.match(r"^\d+\.\s+", nxt)):
                break
            buf.append(nxt)
            i += 1
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(par, " ".join(buf))

    doc.core_properties.title = title
    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2], sys.argv[3])
